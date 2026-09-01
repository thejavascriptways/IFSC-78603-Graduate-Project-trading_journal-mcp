from __future__ import annotations

import math
import os
import subprocess
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "proposal"
DIAGRAM_DIR = OUTPUT_DIR / "diagrams"
DOCX_PATH = OUTPUT_DIR / "Trading_Journal_MCP_Proposal_Revised.docx"
PDF_PATH = OUTPUT_DIR / "Trading_Journal_MCP_Proposal_Revised.pdf"

BLUE = "1F6F9C"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "EAF3FB"
GRAY = "F4F6F9"
DARK_GRAY = "374151"
GREEN = "DDF6E7"
GOLD = "FFF3CD"
RED = "FCE8E8"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt: ImageFont.ImageFont) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = LIGHT_BLUE,
    outline: str = BLUE,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=f"#{fill}", outline=f"#{outline}", width=3)
    title_font = font(25, True)
    body_font = font(21)
    draw.text((x1 + 20, y1 + 18), title, fill=f"#{DARK_BLUE}", font=title_font)
    if body:
        y = y1 + 56
        for line in wrap(draw, body, x2 - x1 - 40, body_font):
            draw.text((x1 + 20, y), line, fill=f"#{DARK_GRAY}", font=body_font)
            y += 28


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = DARK_BLUE) -> None:
    draw.line((start, end), fill=f"#{color}", width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    spread = 0.55
    points = [
        end,
        (
            int(end[0] - length * math.cos(angle - spread)),
            int(end[1] - length * math.sin(angle - spread)),
        ),
        (
            int(end[0] - length * math.cos(angle + spread)),
            int(end[1] - length * math.sin(angle + spread)),
        ),
    ]
    draw.polygon(points, fill=f"#{color}")


def diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1500, 88), fill=f"#{DARK_BLUE}")
    draw.text((45, 25), title, fill="white", font=font(34, True))
    return img, draw


def save_diagram(img: Image.Image, name: str) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    img.save(path, "PNG")
    return path


def make_diagrams() -> dict[str, Path]:
    diagrams: dict[str, Path] = {}

    img, draw = diagram_canvas("Figure 1. Overall System Architecture")
    box(draw, (60, 150, 360, 310), "User", "Uses browser UI or CLI to manage and review portfolio.", GREEN)
    box(draw, (560, 130, 940, 330), "Trading Journal App", "FastAPI web app, JSON APIs, service layer, SQLite database.", LIGHT_BLUE)
    box(draw, (1100, 140, 1440, 300), "Alpaca API", "External market data for stocks, ETFs, and supported options.", GOLD)
    box(draw, (515, 450, 735, 640), "Journal MCP", "Accounts, trades, positions, summaries, prompts.", GRAY)
    box(draw, (780, 450, 1040, 640), "Market Data MCP", "Provider status, quote snapshots, market-data prompt.", GRAY)
    box(draw, (120, 520, 430, 700), "External MCP Client", "CLI demo or future client discovers and calls MCP capabilities.", GREEN)
    arrow(draw, (360, 230), (560, 230))
    arrow(draw, (940, 220), (1100, 220))
    arrow(draw, (635, 330), (635, 450))
    arrow(draw, (910, 450), (910, 330))
    arrow(draw, (430, 610), (515, 560))
    arrow(draw, (1040, 545), (1210, 300))
    diagrams["architecture"] = save_diagram(img, "architecture.png")

    img, draw = diagram_canvas("Figure 2. Day-to-Day User Workflow")
    steps = [
        ("Import Holdings", "Manually enter existing Fidelity positions."),
        ("Enter Trades", "Record buys and sells with required trade reason."),
        ("Refresh Quotes", "Use Market Data MCP to update open-position marks."),
        ("Review P&L", "View realized, unrealized, account, and asset-class performance."),
        ("Journal Review", "Use prompts and trade reasons to analyze decisions."),
    ]
    x = 70
    y = 245
    for i, (t, b) in enumerate(steps):
        box(draw, (x + i * 280, y, x + i * 280 + 230, y + 220), t, b, LIGHT_BLUE)
        if i < len(steps) - 1:
            arrow(draw, (x + i * 280 + 230, y + 110), (x + (i + 1) * 280 - 10, y + 110))
    diagrams["workflow"] = save_diagram(img, "workflow.png")

    img, draw = diagram_canvas("Figure 3. Trade and P&L Lifecycle")
    box(draw, (80, 160, 390, 310), "Opening Holding", "Creates initial position with quantity and average cost.", GREEN)
    box(draw, (590, 150, 920, 320), "Manual Buy", "Increases quantity and recalculates average cost basis.", LIGHT_BLUE)
    box(draw, (1090, 150, 1410, 320), "Manual Sell", "Reduces quantity and calculates realized P&L.", GOLD)
    box(draw, (340, 530, 650, 700), "Open Position", "Market price creates unrealized P&L.", GRAY)
    box(draw, (880, 530, 1190, 700), "Closed Position", "Full sale moves position to closed history.", RED)
    arrow(draw, (390, 235), (590, 235))
    arrow(draw, (920, 235), (1090, 235))
    arrow(draw, (755, 320), (495, 530))
    arrow(draw, (1250, 320), (1035, 530))
    arrow(draw, (650, 615), (880, 615))
    diagrams["trade_lifecycle"] = save_diagram(img, "trade_lifecycle.png")

    img, draw = diagram_canvas("Figure 4. MCP Client-Server Interaction")
    box(draw, (70, 170, 390, 720), "External MCP Client", "1. Initialize session\n2. List tools/resources/prompts\n3. Call selected capability\n4. Read structured response", GREEN)
    box(draw, (600, 150, 930, 350), "Trading Journal MCP", "Portfolio tools, resources, and prompts.", LIGHT_BLUE)
    box(draw, (600, 520, 930, 720), "Market Data MCP", "Quote tools, capabilities resource, and prompt.", LIGHT_BLUE)
    box(draw, (1110, 330, 1430, 540), "Trading Journal App", "Service layer executes operations and returns structured results.", GRAY)
    arrow(draw, (390, 300), (600, 250))
    arrow(draw, (600, 300), (390, 355))
    arrow(draw, (390, 590), (600, 620))
    arrow(draw, (600, 670), (390, 650))
    arrow(draw, (930, 250), (1110, 410))
    arrow(draw, (930, 620), (1110, 460))
    diagrams["mcp_interaction"] = save_diagram(img, "mcp_interaction.png")

    img, draw = diagram_canvas("Figure 5. Market Data Refresh Flow")
    box(draw, (70, 170, 390, 330), "Open Positions", "Symbols and asset classes are collected from current portfolio.", GREEN)
    box(draw, (560, 150, 940, 350), "In-App MCP Client", "Starts MCP session and calls Market Data MCP tools.", LIGHT_BLUE)
    box(draw, (1100, 160, 1430, 330), "Market Data MCP", "Checks provider capability and requests snapshots.", GRAY)
    box(draw, (1100, 550, 1430, 720), "Alpaca", "Returns supported quote snapshots.", GOLD)
    box(draw, (560, 550, 940, 720), "Position Marks", "Market price, market value, and unrealized P&L are updated.", GREEN)
    arrow(draw, (390, 250), (560, 250))
    arrow(draw, (940, 250), (1100, 250))
    arrow(draw, (1265, 330), (1265, 550))
    arrow(draw, (1100, 635), (940, 635))
    arrow(draw, (750, 550), (750, 350))
    diagrams["market_refresh"] = save_diagram(img, "market_refresh.png")

    img, draw = diagram_canvas("Figure 6. Proposed Two-Person Work Plan")
    box(draw, (80, 155, 460, 355), "Team Member 1", "Co-develops app features, MCP workflows, testing, and article sections.", LIGHT_BLUE)
    box(draw, (80, 510, 460, 710), "Team Member 2", "Co-develops app features, MCP workflows, usability review, and article sections.", GREEN)
    box(draw, (620, 160, 930, 350), "Shared Backlog", "Weekly scope review, issue tracking, acceptance criteria.", GRAY)
    box(draw, (1080, 160, 1400, 350), "Prototype Delivery", "Working app, tests, demo scripts, recorded walkthroughs.", GOLD)
    box(draw, (620, 520, 930, 710), "Research Article", "Joint writing: related work, architecture analysis, findings, lessons learned.", GRAY)
    box(draw, (1080, 520, 1400, 710), "Final Submission", "Application + article + demo evidence.", LIGHT_BLUE)
    arrow(draw, (460, 255), (620, 255))
    arrow(draw, (460, 610), (620, 610))
    arrow(draw, (930, 255), (1080, 255))
    arrow(draw, (930, 610), (1080, 610))
    arrow(draw, (780, 350), (780, 520))
    arrow(draw, (1240, 350), (1240, 520))
    diagrams["team_plan"] = save_diagram(img, "team_plan.png")

    return diagrams


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def style_table(table, widths: list[int], header: bool = True) -> None:
    set_table_width(table, widths)
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
            if header and r_idx == 0:
                set_cell_shading(cell, "F4F6F9")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(11, 37, 69)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Trading Journal MCP Proposal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(f" {text}")
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(75, 85, 99)


def add_status_table(doc: Document) -> None:
    rows = [
        ("Area", "Current Prototype", "Planned Full Application"),
        ("Portfolio model", "Accounts, instruments, trades, positions, open and closed position states.", "Broker-aware data model with sync history and production audit metadata."),
        ("Trade journaling", "Manual trades with required reason, notes, fees, and realized P&L on sells.", "Richer review workflow, tagging, attachments, and post-trade reflection prompts."),
        ("Market data", "Alpaca snapshots for supported stocks, ETFs, and options through Market Data MCP.", "Provider hardening, entitlement handling, fallback pricing, and clearer asset-class coverage."),
        ("MCP capability", "Two MCP servers, resources, prompts, internal MCP client, and external CLI client.", "Authenticated remote MCP access, browser MCP console, and stronger consent controls."),
        ("Testing", "Automated tests for trade flow, P&L, dashboard summaries, and MCP behavior.", "Expanded integration tests, security checks, and deployment smoke tests."),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    style_table(table, [1900, 3830, 3630])


def add_scope_table(doc: Document) -> None:
    rows = [
        ("Phase", "Primary Deliverables", "Feasibility Rationale"),
        ("Prototype validation", "Manual holdings, trade entry, P&L, two MCP servers, CLI client.", "Already completed as preliminary work."),
        ("Core full application", "Improved UI, richer analytics, broker sync foundation, secure local config.", "Builds directly on working service layer and tests."),
        ("MCP maturity", "MCP console, authenticated remote access plan, stronger prompt/resource design.", "Incremental extension of existing MCP endpoints."),
        ("Research article", "Related work, architecture analysis, security discussion, lessons learned.", "Uses the prototype as evidence and case study material."),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    style_table(table, [1650, 4200, 3510])


def add_references(doc: Document) -> None:
    refs = [
        ("Model Context Protocol. (2026). What is the Model Context Protocol (MCP)?", "https://modelcontextprotocol.io/docs/getting-started/intro"),
        ("Model Context Protocol. (2025). Specification, 2025-06-18.", "https://modelcontextprotocol.io/specification/2025-06-18"),
        ("Soria Parra, D. (2026). The 2026 MCP Roadmap. Model Context Protocol Blog.", "https://blog.modelcontextprotocol.io/posts/2026-mcp-roadmap/"),
        ("Eleventh Hour Enthusiast. (2025). Model Context Protocol (MCP): Landscape, Security Threats, and Future Research Directions.", "https://medium.com/@EleventhHourEnthusiast/model-context-protocol-mcp-landscape-security-threats-and-future-research-directions-488b8d2eade8"),
        ("Decode the Future. (2026). What Is MCP? Model Context Protocol Explained for 2026.", "https://decodethefuture.org/en/what-is-mcp-model-context-protocol/"),
        ("FastMCP Documentation. Getting Started.", "https://gofastmcp.com/getting-started/welcome"),
        ("FastAPI Documentation.", "https://fastapi.tiangolo.com/"),
        ("SQLAlchemy Documentation.", "https://docs.sqlalchemy.org/"),
        ("Alpaca Documentation. Getting Started.", "https://docs.alpaca.markets/us/docs/getting-started"),
        ("Alpaca Documentation. About Market Data API.", "https://docs.alpaca.markets/us/docs/about-market-data-api"),
        ("Alpaca Market Data Overview.", "https://alpaca.markets/data"),
    ]
    for idx, (label, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(f"[{idx}] {label} ")
        add_hyperlink(p, url, url)


def build_docx(diagrams: dict[str, Path]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    doc.add_paragraph("Project Proposal", style="Subtitle")
    doc.add_paragraph("Trading Journal: Leveraging MCP in the Financial Domain", style="Title")
    doc.add_paragraph("Revised proposal with related work, preliminary prototype evidence, project coordination plan, and application flowcharts.", style="Subtitle")

    add_callout(
        doc,
        "Project thesis.",
        "Trading Journal is proposed as a Python-based portfolio and trade-journaling application that also serves as a practical case study for Model Context Protocol client-server design in a financial software setting.",
    )

    doc.add_heading("Project Team", level=1)
    add_paragraph(
        doc,
        "This project will be completed by a two-person team. Both team members will contribute to the technical implementation and the research/article components of the project. The responsibilities are intentionally shared so each team member develops hands-on experience with MCP, application design, financial-domain workflows, testing, documentation, and research communication.",
    )
    add_bullets(
        doc,
        [
            "Team Member 1: Deepak Singla. Co-develop the Trading Journal application, implement and test MCP server/client workflows, contribute to portfolio and market-data features, participate in related-work analysis, and help write the final article.",
            "Team Member 2: Ram Kiran Hota. Co-develop application features, assist with MCP integration and testing, contribute to UI/workflow design, participate in related-work analysis, and help write the final article.",
            "Shared responsibilities: weekly planning, implementation reviews, prototype testing, documentation, demo preparation, architecture diagrams, security discussion, and final submission.",
        ],
    )

    doc.add_heading("Introduction", level=1)
    add_paragraph(
        doc,
        "Model Context Protocol (MCP) is an open standard for connecting AI applications to external systems, including data sources, tools, and reusable workflows. The official MCP documentation describes MCP as a standard that lets applications such as Claude or ChatGPT connect to external files, databases, tools, and prompts through a consistent interface [1]. In practical terms, MCP gives an AI-capable application a way to discover what a system can do, call structured tools, read resources, and use reusable prompts without building a separate custom integration for every data source.",
    )
    add_paragraph(
        doc,
        "This project proposes the continued development of Trading Journal, a Python-based portfolio tracking and trade-journaling application. The application is inspired by the day-to-day usability of retail brokerage portfolio views, but it adds a journaling layer that records the reason for each trade. The project has two connected goals. The first goal is technical and academic: to learn and demonstrate MCP by building a real application that exposes MCP servers and also consumes MCP capabilities as a client. The second goal is practical: to build a useful local-first application for tracking holdings, entering trades, reviewing profit and loss, and analyzing investment decisions.",
    )
    add_paragraph(
        doc,
        "The preliminary prototype already demonstrates the critical path for the project. It supports manual holding import, manual trade entry, realized and unrealized profit/loss calculations, live market-data integration for supported instruments, two MCP servers, and an external command-line MCP client. This makes the proposal stronger than a purely conceptual plan because the most important architecture choices have already been tested in a working system.",
    )

    doc.add_heading("Summary of the Proposal", level=1)
    add_paragraph(
        doc,
        "The proposed system is a full-stack Python application for portfolio management, trade journaling, and MCP-based interoperability. In its first complete version, the application will allow users to import current holdings manually, record new buy and sell trades, close positions, and review performance across open and closed positions. Every trade will require a reason field so the user can later analyze not only what happened financially, but also why the decision was made.",
    )
    add_paragraph(
        doc,
        "The project will also serve as a learning-oriented MCP implementation. The Trading Journal application will expose portfolio and journaling features through a Trading Journal MCP server. Market-data features will be exposed through a separate Market Data MCP server. A command-line MCP client has already been created to demonstrate how an outside client can discover and use those capabilities. This gives the project a clear academic contribution: it shows MCP in a realistic domain rather than as a small artificial example.",
    )
    add_paragraph(
        doc,
        "The final project will balance application development with a high-quality submitted article. The article will discuss the architecture, related work, security considerations, implementation decisions, and lessons learned from building a financial MCP application. The application and article together form a reasonable two-person project because both team members will contribute to both facets, with specific ownership rotating by milestone rather than by a strict developer/writer split.",
    )

    doc.add_heading("Related Work", level=1)
    add_paragraph(
        doc,
        "The official MCP documentation frames MCP as an open protocol for connecting AI applications to external systems [1]. This is directly relevant to Trading Journal because the project treats the portfolio application as a provider of structured capabilities rather than only as a web interface. The application exposes tools such as portfolio summary retrieval and trade entry, resources such as position snapshots, and prompts such as daily portfolio reviews. This follows the core MCP model of making external context and capabilities available to AI systems through a standard interface.",
    )
    add_paragraph(
        doc,
        "The MCP specification describes the protocol as a standardized way for applications to share contextual information, expose tools, and build composable workflows using JSON-RPC 2.0, capability negotiation, resources, prompts, and tools [2]. The specification is important for this project because it provides the conceptual boundary between a normal REST API and an MCP-enabled application. In Trading Journal, the MCP layer is not only another endpoint; it is a capability surface that can be discovered and used by external clients.",
    )
    add_paragraph(
        doc,
        "The 2026 MCP roadmap identifies several areas that are especially relevant to this project: transport scalability, agent communication, governance maturation, enterprise readiness, stronger metadata discovery, and deeper security and authorization work [3]. The current prototype is intentionally local-first, but the roadmap helps shape the future direction of the project. For example, the prototype uses streamable HTTP locally, while the full version will need to consider authentication, deployment, authorization, and remote-server readiness before it can safely be used over the internet.",
    )
    add_paragraph(
        doc,
        "The MCP landscape and security review recommended by Dr. Batthula is especially important because financial applications involve sensitive user data and potentially high-impact actions [4]. That work highlights risks across the MCP server lifecycle, including malicious or misleading tools, sandbox escape risks, configuration drift, privilege persistence, lack of centralized oversight, and authentication gaps. These concerns directly influence this proposal. The full application should not expose unrestricted trade or account actions to remote clients without authentication, user consent, careful logging, and least-privilege tool design.",
    )
    add_paragraph(
        doc,
        "The Decode the Future overview provides a useful applied explanation of MCP primitives: tools act, resources provide read-only context, and prompts standardize reusable workflows [5]. This distinction maps cleanly onto the Trading Journal prototype. Adding a manual trade is a tool because it changes application state. Reading portfolio positions is a resource because it retrieves data without changing it. Generating a portfolio review instruction is a prompt because it gives an AI assistant a reusable analysis template.",
    )
    add_paragraph(
        doc,
        "FastAPI is used as the application foundation because it is a modern Python web framework designed for API development with type hints, automatic documentation, and production-oriented features [7]. SQLAlchemy supports the database model and service-layer persistence [8]. Alpaca is used as the preliminary market-data provider because its Market Data API supports HTTP and WebSocket access to historical and real-time market data, while also making clear that feed coverage depends on subscription level and asset class [10]. These choices keep the prototype practical while preserving room for future provider or broker integration.",
    )

    doc.add_heading("Proposed Application Scope", level=1)
    add_paragraph(
        doc,
        "The full application will focus on the intersection of portfolio tracking, trade journaling, and MCP-based integration. The intended first complete version will remain scoped enough for a two-person team, while still being substantial enough to demonstrate meaningful software engineering and research value.",
    )
    add_bullets(
        doc,
        [
            "Support manual import of existing holdings and manual entry of buy and sell trades.",
            "Track open and closed positions using average-cost accounting.",
            "Calculate realized P&L, unrealized P&L, market value, cost basis, and P&L percentages.",
            "Report performance overall, by account, and by account plus asset class.",
            "Require a reason for each trade so the application functions as both a tracker and a decision journal.",
            "Use external market data for supported stocks, ETFs, and options, while clearly documenting limitations for mutual funds, bonds, and options entitlements.",
            "Expose portfolio capabilities through a Trading Journal MCP server.",
            "Expose market-data capabilities through a separate Market Data MCP server.",
            "Provide an external command-line MCP client and later a browser-based MCP inspection console.",
            "Prepare the application for later deployment with authentication, authorization, and safer remote access.",
        ],
    )
    add_scope_table(doc)

    doc.add_heading("Preliminary Work Completed", level=1)
    add_paragraph(
        doc,
        "A working prototype has already been developed to validate the main technical direction of the project. The prototype is built using Python, FastAPI, SQLAlchemy, SQLite, Jinja templates, and MCP. The purpose of this early work was not to complete the production application, but to prove that the core ideas are feasible: portfolio tracking, trade journaling, profit/loss calculation, live market-data integration, and MCP-based communication between clients and servers.",
    )
    add_paragraph(
        doc,
        "The prototype currently runs as a local-first web application. It includes a database-backed portfolio model with accounts, instruments, trades, and positions. Two default account types have been created: an IBKR Live account to represent a future broker-connected account and a Manual Fidelity account to support manual import of holdings. This allows the prototype to support the initial use case where existing Fidelity holdings can be entered manually before automated brokerage synchronization is added later.",
    )
    add_paragraph(
        doc,
        "A manual holding import workflow has been implemented. This allows a user to enter existing positions into the system with account, symbol, asset class, quantity, average cost, opening date, description, and notes. After the initial holdings are entered, the user can maintain the portfolio by manually entering new buy and sell trades. The trade-entry workflow includes account, symbol, asset class, trade date, side, quantity, price, fees, currency, notes, and a required reason for trade.",
    )
    add_paragraph(
        doc,
        "Position tracking and P&L calculations have also been implemented. The application maintains open positions using average-cost accounting. Buy trades increase position quantity and update average cost. Sell trades reduce quantity and calculate realized P&L. If a sell trade closes the full position, the position is moved from open positions to closed position history. Unrealized P&L is calculated for open positions using current market prices.",
    )
    add_paragraph(
        doc,
        "The current web interface includes pages for the dashboard, trades, open positions, closed positions, opening holding import, manual trade entry, and market data. The dashboard reports overall performance, account-level performance, and account plus asset-class performance. The market-data page displays quote information for open and closed positions where supported. The positions page includes a quote refresh workflow that updates open-position marks through MCP.",
    )
    add_paragraph(
        doc,
        "The preliminary work also includes two MCP servers. The Trading Journal MCP server exposes tools for portfolio summary retrieval, account listing, position listing, trade listing, opening holding creation, and manual trade creation. It also exposes resources such as portfolio://summary and portfolio://positions, and prompts such as daily_portfolio_review and journal_follow_up. The Market Data MCP server exposes tools for provider capability checks, equity snapshots, and option snapshots, along with the market-data://capabilities resource and a market_data_health_check prompt.",
    )
    add_paragraph(
        doc,
        "The application also acts as an MCP client internally. When the user refreshes market prices, the Trading Journal application starts an MCP client session and calls the Market Data MCP server rather than directly embedding market-data calls in the page route. This demonstrates both sides of MCP in one system: the application exposes capabilities as a server and consumes capabilities as a client.",
    )
    add_paragraph(
        doc,
        "A separate command-line MCP client has been created to support demonstration and evaluation. This client can explain MCP concepts, list available MCP servers, discover tools/resources/prompts, call tools, read resources, render prompts, and run guided workflows such as portfolio-review, market-check, and client-demo. This is important because it proves that another process can use the application's MCP capabilities without using the browser interface or importing the application source code.",
    )
    add_status_table(doc)

    doc.add_heading("Application Flowcharts", level=1)
    add_paragraph(
        doc,
        "The following figures summarize the current prototype and proposed full application architecture. They are included to make the project easier to review visually and to clarify how MCP fits into the application.",
    )
    add_figure(doc, diagrams["architecture"], "Figure 1. The application combines a local FastAPI web app, database-backed portfolio logic, two MCP servers, an external MCP client, and an external market-data provider.")
    add_figure(doc, diagrams["workflow"], "Figure 2. Day-to-day usage begins with importing holdings, then entering trades, refreshing quotes, reviewing P&L, and analyzing journal decisions.")
    add_figure(doc, diagrams["trade_lifecycle"], "Figure 3. Trade activity updates positions and moves fully sold positions into closed history while preserving realized P&L.")
    add_figure(doc, diagrams["mcp_interaction"], "Figure 4. An external MCP client discovers and uses tools, resources, and prompts from either MCP server.")
    add_figure(doc, diagrams["market_refresh"], "Figure 5. Market-data refresh uses an in-app MCP client to call the Market Data MCP server and update open-position marks.")
    add_figure(doc, diagrams["team_plan"], "Figure 6. A two-person structure keeps both members involved in implementation and research/article work while maintaining a shared backlog and final integrated deliverable.")

    doc.add_heading("Two-Person Project Plan and Scope Justification", level=1)
    add_paragraph(
        doc,
        "The project is reasonable for a two-person team because the preliminary prototype has already reduced the highest technical uncertainty. The remaining work is not starting from zero; it is an extension and hardening of an existing codebase. The work will not be split into separate technical and writing roles. Instead, both team members will contribute to both the application and the article, with individual task ownership rotating by milestone.",
    )
    add_paragraph(
        doc,
        "For example, one milestone may have Team Member 1 leading a backend MCP feature while Team Member 2 leads the matching UI/demo documentation. In a later milestone, those roles can rotate so both members gain experience across implementation, testing, usability, related work, and article writing. Both members will review each other's work through a shared backlog, weekly check-ins, and milestone-based demos.",
    )
    add_paragraph(
        doc,
        "The project scope is justified because the final deliverable is not only a production-ready trading platform. It is a technical project plus a high-quality submitted article. That means the team can focus on a strong, demonstrable subset of features: reliable manual portfolio tracking, correct P&L logic, meaningful MCP integration, clear market-data boundaries, and a well-written analysis of the architecture and lessons learned. More advanced features such as full internet deployment, automated broker trading, and production-grade multi-user support can be treated as future work unless time permits.",
    )
    add_bullets(
        doc,
        [
            "Weekly coordination: one planning meeting, one demo/checkpoint, and shared issue tracking.",
            "Rotating technical ownership: MCP architecture, backend services, data model, broker-sync foundation, automated tests.",
            "Rotating research and UX ownership: related work, proposal/article writing, diagrams, usability review, demo script, documentation.",
            "Shared responsibilities: scope decisions, security review, final integration, presentation, and submitted article quality.",
        ],
    )

    doc.add_heading("Security and Ethics Considerations", level=1)
    add_paragraph(
        doc,
        "Security is central to this project because the financial domain involves sensitive account data and potentially high-impact actions. The prototype is intentionally local-first, and the current MCP endpoints are intended for local development and demonstration. A full internet-deployed version must add authentication, authorization, transport security, logging, and explicit user-consent flows before exposing MCP capabilities remotely.",
    )
    add_paragraph(
        doc,
        "The related work on MCP security identifies risks such as tool poisoning, ambiguous tool names, sandbox escape, configuration drift, and authorization gaps [4]. The project will address these risks by keeping tool names explicit, separating read-only resources from write-capable tools, documenting tool side effects, avoiding automatic execution of sensitive actions, and designing future remote access around least privilege and user approval. In a financial application, MCP tools that change portfolio state must be treated differently from resources that only read summary data.",
    )
    add_paragraph(
        doc,
        "The project is intended for journaling and analysis, not for automated trading execution in the first version. This boundary reduces risk and keeps the application aligned with the academic objective: to study MCP architecture and build a useful portfolio journal. Any future integration with live brokerage trading would require additional compliance, risk, and user-confirmation controls.",
    )

    doc.add_heading("Evaluation Plan", level=1)
    add_paragraph(
        doc,
        "The project will be evaluated through both technical and usability evidence. Technical evaluation will focus on correctness of trade workflows, P&L calculations, market-data behavior, MCP discovery, MCP tool calls, resource reads, prompt rendering, and regression tests. Usability evaluation will focus on whether the application supports the intended daily workflow: importing holdings, entering trades, recording trade reasons, closing positions, reviewing performance, and demonstrating MCP behavior clearly.",
    )
    add_bullets(
        doc,
        [
            "Functional tests for trade entry, holding import, P&L calculations, and closed-position behavior.",
            "MCP tests for tool discovery, resource access, prompt rendering, and multi-server workflows.",
            "Demo scripts showing CLI-based external client behavior.",
            "Manual usability checks for dashboard, positions, trades, market data, and journaling workflows.",
            "Article-based evaluation discussing architecture tradeoffs, related work, and future improvements.",
        ],
    )

    doc.add_heading("Planned Next Steps", level=1)
    add_numbered(
        doc,
        [
            "Complete the browser-based MCP console so MCP tools, resources, and prompts can be inspected without relying only on terminal commands.",
            "Improve UI polish and navigation for daily portfolio use.",
            "Expand market-data handling, including clearer unsupported-asset messages and provider configuration guidance.",
            "Prepare an IBKR-first synchronization design while keeping manual Fidelity import available.",
            "Add authentication and deployment planning for future internet access.",
            "Develop the submitted article using the prototype as a case study in MCP architecture for financial applications.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "Trading Journal is a feasible and meaningful project because it combines a practical user-facing application with a current technical research topic. The preliminary prototype already proves the central idea: a portfolio journal can expose useful capabilities through MCP, consume market-data capabilities through another MCP server, and be demonstrated by an external MCP client. The proposed next phase will mature this prototype into a stronger application and a well-supported article that discusses MCP's role, benefits, limitations, and security considerations in the financial domain.",
    )

    doc.add_heading("References", level=1)
    add_references(doc)

    doc.save(DOCX_PATH)


def render_docx() -> None:
    render_dir = OUTPUT_DIR / "rendered"
    render_dir.mkdir(parents=True, exist_ok=True)
    renderer = Path("/Users/deepaksingla/.codex/plugins/cache/openai-primary-runtime/documents/26.727.11326/skills/documents/render_docx.py")
    subprocess.run(
        [
            "/Users/deepaksingla/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
            str(renderer),
            str(DOCX_PATH),
            "--output_dir",
            str(render_dir),
            "--emit_pdf",
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    emitted_pdf = render_dir / f"{DOCX_PATH.stem}.pdf"
    if emitted_pdf.exists():
        PDF_PATH.write_bytes(emitted_pdf.read_bytes())


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ProposalTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=29,
            textColor=colors.HexColor(f"#{DARK_BLUE}"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "ProposalSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            textColor=colors.HexColor(f"#{DARK_GRAY}"),
            alignment=TA_CENTER,
            spaceAfter=20,
        ),
        "h1": ParagraphStyle(
            "Heading1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=16,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "Heading2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor(f"#1F4D78"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.3,
            leading=14.2,
            alignment=TA_JUSTIFY,
            spaceAfter=7,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8.7,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=10,
        ),
        "callout": ParagraphStyle(
            "Callout",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13.5,
            textColor=colors.HexColor(f"#{DARK_BLUE}"),
            leftIndent=8,
            rightIndent=8,
            spaceAfter=8,
        ),
        "table": ParagraphStyle(
            "Table",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.1,
            leading=10.2,
            spaceAfter=0,
        ),
        "ref": ParagraphStyle(
            "References",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.5,
            leading=9.0,
            leftIndent=18,
            firstLineIndent=-18,
            spaceAfter=3,
        ),
    }


def pp(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def bullet_list(items: list[str], styles: dict[str, ParagraphStyle]) -> ListFlowable:
    return ListFlowable(
        [ListItem(pp(item, styles["body"]), leftIndent=12) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
        bulletFontSize=8,
    )


def number_list(items: list[str], styles: dict[str, ParagraphStyle]) -> ListFlowable:
    return ListFlowable(
        [ListItem(pp(item, styles["body"]), leftIndent=12) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def proposal_table(rows: list[tuple[str, ...]], widths: list[float], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[Paragraph(escape(cell), styles["table"]) for cell in row] for row in rows]
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{DARK_BLUE}")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def add_pdf_figure(story: list, path: Path, caption: str, styles: dict[str, ParagraphStyle]) -> None:
    story.append(KeepTogether([RLImage(str(path), width=6.25 * inch, height=3.75 * inch), pp(caption, styles["caption"])]))


def build_pdf(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.72 * inch,
        title="Trading Journal MCP Proposal Revised",
    )
    story: list = []

    story.append(pp("Project Proposal", styles["subtitle"]))
    story.append(pp("Trading Journal: Leveraging MCP in the Financial Domain", styles["title"]))
    story.append(pp("Revised proposal with related work, preliminary prototype evidence, project coordination plan, and application flowcharts.", styles["subtitle"]))
    callout = Table(
        [[Paragraph("<b>Project thesis.</b> Trading Journal is proposed as a Python-based portfolio and trade-journaling application that also serves as a practical case study for Model Context Protocol client-server design in a financial software setting.", styles["callout"])]],
        colWidths=[6.4 * inch],
        hAlign="LEFT",
    )
    callout.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#CBD5E1")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([callout, Spacer(1, 8)])

    story.append(pp("Project Team", styles["h1"]))
    story.append(pp("This project will be completed by a two-person team. Both team members will contribute to the technical implementation and the research/article components of the project. The responsibilities are intentionally shared so each team member develops hands-on experience with MCP, application design, financial-domain workflows, testing, documentation, and research communication.", styles["body"]))
    story.append(
        bullet_list(
            [
                "Team Member 1: Deepak Singla. Co-develop the Trading Journal application, implement and test MCP server/client workflows, contribute to portfolio and market-data features, participate in related-work analysis, and help write the final article.",
                "Team Member 2: Ram Kiran Hota. Co-develop application features, assist with MCP integration and testing, contribute to UI/workflow design, participate in related-work analysis, and help write the final article.",
                "Shared responsibilities: weekly planning, implementation reviews, prototype testing, documentation, demo preparation, architecture diagrams, security discussion, and final submission.",
            ],
            styles,
        )
    )

    sections = [
        (
            "Introduction",
            [
                "Model Context Protocol (MCP) is an open standard for connecting AI applications to external systems, including data sources, tools, and reusable workflows. The official MCP documentation describes MCP as a standard that lets applications such as Claude or ChatGPT connect to external files, databases, tools, and prompts through a consistent interface [1]. In practical terms, MCP gives an AI-capable application a way to discover what a system can do, call structured tools, read resources, and use reusable prompts without building a separate custom integration for every data source.",
                "This project proposes the continued development of Trading Journal, a Python-based portfolio tracking and trade-journaling application. The application is inspired by the day-to-day usability of retail brokerage portfolio views, but it adds a journaling layer that records the reason for each trade. The project has two connected goals. The first goal is technical and academic: to learn and demonstrate MCP by building a real application that exposes MCP servers and also consumes MCP capabilities as a client. The second goal is practical: to build a useful local-first application for tracking holdings, entering trades, reviewing profit and loss, and analyzing investment decisions.",
                "The preliminary prototype already demonstrates the critical path for the project. It supports manual holding import, manual trade entry, realized and unrealized profit/loss calculations, live market-data integration for supported instruments, two MCP servers, and an external command-line MCP client. This makes the proposal stronger than a purely conceptual plan because the most important architecture choices have already been tested in a working system.",
            ],
        ),
        (
            "Summary of the Proposal",
            [
                "The proposed system is a full-stack Python application for portfolio management, trade journaling, and MCP-based interoperability. In its first complete version, the application will allow users to import current holdings manually, record new buy and sell trades, close positions, and review performance across open and closed positions. Every trade will require a reason field so the user can later analyze not only what happened financially, but also why the decision was made.",
                "The project will also serve as a learning-oriented MCP implementation. The Trading Journal application will expose portfolio and journaling features through a Trading Journal MCP server. Market-data features will be exposed through a separate Market Data MCP server. A command-line MCP client has already been created to demonstrate how an outside client can discover and use those capabilities. This gives the project a clear academic contribution: it shows MCP in a realistic domain rather than as a small artificial example.",
                "The final project will balance application development with a high-quality submitted article. The article will discuss the architecture, related work, security considerations, implementation decisions, and lessons learned from building a financial MCP application. The application and article together form a reasonable two-person project because both team members will contribute to both facets, with specific ownership rotating by milestone rather than by a strict developer/writer split.",
            ],
        ),
        (
            "Related Work",
            [
                "The official MCP documentation frames MCP as an open protocol for connecting AI applications to external systems [1]. This is directly relevant to Trading Journal because the project treats the portfolio application as a provider of structured capabilities rather than only as a web interface. The application exposes tools such as portfolio summary retrieval and trade entry, resources such as position snapshots, and prompts such as daily portfolio reviews.",
                "The MCP specification describes the protocol as a standardized way for applications to share contextual information, expose tools, and build composable workflows using JSON-RPC 2.0, capability negotiation, resources, prompts, and tools [2]. The specification is important for this project because it provides the conceptual boundary between a normal REST API and an MCP-enabled application.",
                "The 2026 MCP roadmap identifies several areas that are especially relevant to this project: transport scalability, agent communication, governance maturation, enterprise readiness, stronger metadata discovery, and deeper security and authorization work [3]. The current prototype is intentionally local-first, but the roadmap helps shape the future direction of the project.",
                "The MCP landscape and security review recommended by Dr. Batthula is especially important because financial applications involve sensitive user data and potentially high-impact actions [4]. That work highlights risks across the MCP server lifecycle, including malicious or misleading tools, sandbox escape risks, configuration drift, privilege persistence, lack of centralized oversight, and authentication gaps.",
                "The Decode the Future overview provides a useful applied explanation of MCP primitives: tools act, resources provide read-only context, and prompts standardize reusable workflows [5]. This distinction maps cleanly onto the Trading Journal prototype. Adding a manual trade is a tool because it changes application state. Reading portfolio positions is a resource because it retrieves data without changing it.",
                "FastAPI is used as the application foundation because it is a modern Python web framework designed for API development with type hints, automatic documentation, and production-oriented features [7]. SQLAlchemy supports the database model and service-layer persistence [8]. Alpaca is used as the preliminary market-data provider because its Market Data API supports HTTP and WebSocket access to historical and real-time market data, while also making clear that feed coverage depends on subscription level and asset class [10].",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        story.append(pp(heading, styles["h1"]))
        for text in paragraphs:
            story.append(pp(text, styles["body"]))

    story.append(pp("Proposed Application Scope", styles["h1"]))
    story.append(pp("The full application will focus on the intersection of portfolio tracking, trade journaling, and MCP-based integration. The intended first complete version will remain scoped enough for a two-person team, while still being substantial enough to demonstrate meaningful software engineering and research value.", styles["body"]))
    story.append(
        bullet_list(
            [
                "Support manual import of existing holdings and manual entry of buy and sell trades.",
                "Track open and closed positions using average-cost accounting.",
                "Calculate realized P&L, unrealized P&L, market value, cost basis, and P&L percentages.",
                "Report performance overall, by account, and by account plus asset class.",
                "Require a reason for each trade so the application functions as both a tracker and a decision journal.",
                "Expose portfolio capabilities through a Trading Journal MCP server and market-data capabilities through a separate Market Data MCP server.",
                "Prepare the application for later deployment with authentication, authorization, and safer remote access.",
            ],
            styles,
        )
    )
    scope_rows = [
        ("Phase", "Primary Deliverables", "Feasibility Rationale"),
        ("Prototype validation", "Manual holdings, trade entry, P&L, two MCP servers, CLI client.", "Already completed as preliminary work."),
        ("Core full application", "Improved UI, richer analytics, broker sync foundation, secure local config.", "Builds directly on working service layer and tests."),
        ("MCP maturity", "MCP console, authenticated remote access plan, stronger prompt/resource design.", "Incremental extension of existing MCP endpoints."),
        ("Research article", "Related work, architecture analysis, security discussion, lessons learned.", "Uses the prototype as evidence and case study material."),
    ]
    story.extend([proposal_table(scope_rows, [1.25 * inch, 3.0 * inch, 2.25 * inch], styles), Spacer(1, 8)])

    story.append(pp("Preliminary Work Completed", styles["h1"]))
    for text in [
        "A working prototype has already been developed to validate the main technical direction of the project. The prototype is built using Python, FastAPI, SQLAlchemy, SQLite, Jinja templates, and MCP. The purpose of this early work was not to complete the production application, but to prove that the core ideas are feasible: portfolio tracking, trade journaling, profit/loss calculation, live market-data integration, and MCP-based communication between clients and servers.",
        "The prototype currently runs as a local-first web application. It includes a database-backed portfolio model with accounts, instruments, trades, and positions. Two default account types have been created: an IBKR Live account to represent a future broker-connected account and a Manual Fidelity account to support manual import of holdings.",
        "A manual holding import workflow has been implemented. This allows a user to enter existing positions into the system with account, symbol, asset class, quantity, average cost, opening date, description, and notes. The trade-entry workflow includes a required reason for trade, making the application a journal as well as a tracker.",
        "Position tracking and P&L calculations have also been implemented. The application maintains open positions using average-cost accounting. Buy trades increase position quantity and update average cost. Sell trades reduce quantity and calculate realized P&L. If a sell trade closes the full position, the position is moved from open positions to closed position history.",
        "The current web interface includes pages for the dashboard, trades, open positions, closed positions, opening holding import, manual trade entry, and market data. The dashboard reports overall performance, account-level performance, and account plus asset-class performance.",
        "The preliminary work also includes two MCP servers. The Trading Journal MCP server exposes tools for portfolio summary retrieval, account listing, position listing, trade listing, opening holding creation, and manual trade creation. It also exposes resources such as portfolio://summary and portfolio://positions, and prompts such as daily_portfolio_review and journal_follow_up.",
        "The Market Data MCP server exposes tools for provider capability checks, equity snapshots, and option snapshots, along with the market-data://capabilities resource and a market_data_health_check prompt. The application also acts as an MCP client internally when refreshing market prices.",
        "A separate command-line MCP client has been created to support demonstration and evaluation. This client can explain MCP concepts, list available MCP servers, discover tools/resources/prompts, call tools, read resources, render prompts, and run guided workflows such as portfolio-review, market-check, and client-demo.",
    ]:
        story.append(pp(text, styles["body"]))
    status_rows = [
        ("Area", "Current Prototype", "Planned Full Application"),
        ("Portfolio model", "Accounts, instruments, trades, positions, open and closed position states.", "Broker-aware data model with sync history and production audit metadata."),
        ("Trade journaling", "Manual trades with required reason, notes, fees, and realized P&L on sells.", "Richer review workflow, tagging, attachments, and post-trade reflection prompts."),
        ("Market data", "Alpaca snapshots for supported stocks, ETFs, and options through Market Data MCP.", "Provider hardening, entitlement handling, fallback pricing, and clearer asset-class coverage."),
        ("MCP capability", "Two MCP servers, resources, prompts, internal MCP client, and external CLI client.", "Authenticated remote MCP access, browser MCP console, and stronger consent controls."),
        ("Testing", "Automated tests for trade flow, P&L, dashboard summaries, and MCP behavior.", "Expanded integration tests, security checks, and deployment smoke tests."),
    ]
    story.extend([proposal_table(status_rows, [1.15 * inch, 2.8 * inch, 2.55 * inch], styles), Spacer(1, 8)])

    story.append(PageBreak())
    story.append(pp("Application Flowcharts", styles["h1"]))
    story.append(pp("The following figures summarize the current prototype and proposed full application architecture. They are included to make the project easier to review visually and to clarify how MCP fits into the application.", styles["body"]))
    add_pdf_figure(story, diagrams["architecture"], "Figure 1. The application combines a local FastAPI web app, database-backed portfolio logic, two MCP servers, an external MCP client, and an external market-data provider.", styles)
    add_pdf_figure(story, diagrams["workflow"], "Figure 2. Day-to-day usage begins with importing holdings, then entering trades, refreshing quotes, reviewing P&L, and analyzing journal decisions.", styles)
    add_pdf_figure(story, diagrams["trade_lifecycle"], "Figure 3. Trade activity updates positions and moves fully sold positions into closed history while preserving realized P&L.", styles)
    add_pdf_figure(story, diagrams["mcp_interaction"], "Figure 4. An external MCP client discovers and uses tools, resources, and prompts from either MCP server.", styles)
    add_pdf_figure(story, diagrams["market_refresh"], "Figure 5. Market-data refresh uses an in-app MCP client to call the Market Data MCP server and update open-position marks.", styles)
    add_pdf_figure(story, diagrams["team_plan"], "Figure 6. A two-person structure keeps both members involved in implementation and research/article work while maintaining a shared backlog and final integrated deliverable.", styles)

    story.append(pp("Two-Person Project Plan and Scope Justification", styles["h1"]))
    for text in [
        "The project is reasonable for a two-person team because the preliminary prototype has already reduced the highest technical uncertainty. The remaining work is not starting from zero; it is an extension and hardening of an existing codebase. The work will not be split into separate technical and writing roles. Instead, both team members will contribute to both the application and the article, with individual task ownership rotating by milestone.",
        "For example, one milestone may have Team Member 1 leading a backend MCP feature while Team Member 2 leads the matching UI/demo documentation. In a later milestone, those roles can rotate so both members gain experience across implementation, testing, usability, related work, and article writing.",
        "The project scope is justified because the final deliverable is not only a production-ready trading platform. It is a technical project plus a high-quality submitted article. That means the team can focus on a strong, demonstrable subset of features: reliable manual portfolio tracking, correct P&L logic, meaningful MCP integration, clear market-data boundaries, and a well-written analysis of the architecture and lessons learned.",
    ]:
        story.append(pp(text, styles["body"]))
    story.append(
        bullet_list(
            [
                "Weekly coordination: one planning meeting, one demo/checkpoint, and shared issue tracking.",
                "Rotating technical ownership: MCP architecture, backend services, data model, broker-sync foundation, automated tests.",
                "Rotating research and UX ownership: related work, proposal/article writing, diagrams, usability review, demo script, documentation.",
                "Shared responsibilities: scope decisions, security review, final integration, presentation, and submitted article quality.",
            ],
            styles,
        )
    )

    story.append(pp("Security and Ethics Considerations", styles["h1"]))
    for text in [
        "Security is central to this project because the financial domain involves sensitive account data and potentially high-impact actions. The prototype is intentionally local-first, and the current MCP endpoints are intended for local development and demonstration. A full internet-deployed version must add authentication, authorization, transport security, logging, and explicit user-consent flows before exposing MCP capabilities remotely.",
        "The related work on MCP security identifies risks such as tool poisoning, ambiguous tool names, sandbox escape, configuration drift, and authorization gaps [4]. The project will address these risks by keeping tool names explicit, separating read-only resources from write-capable tools, documenting tool side effects, avoiding automatic execution of sensitive actions, and designing future remote access around least privilege and user approval.",
        "The project is intended for journaling and analysis, not for automated trading execution in the first version. This boundary reduces risk and keeps the application aligned with the academic objective: to study MCP architecture and build a useful portfolio journal.",
    ]:
        story.append(pp(text, styles["body"]))

    story.append(pp("Evaluation Plan", styles["h1"]))
    story.append(pp("The project will be evaluated through both technical and usability evidence. Technical evaluation will focus on correctness of trade workflows, P&L calculations, market-data behavior, MCP discovery, MCP tool calls, resource reads, prompt rendering, and regression tests. Usability evaluation will focus on whether the application supports the intended daily workflow.", styles["body"]))
    story.append(
        bullet_list(
            [
                "Functional tests for trade entry, holding import, P&L calculations, and closed-position behavior.",
                "MCP tests for tool discovery, resource access, prompt rendering, and multi-server workflows.",
                "Demo scripts showing CLI-based external client behavior.",
                "Manual usability checks for dashboard, positions, trades, market data, and journaling workflows.",
                "Article-based evaluation discussing architecture tradeoffs, related work, and future improvements.",
            ],
            styles,
        )
    )

    story.append(pp("Planned Next Steps", styles["h1"]))
    story.append(
        number_list(
            [
                "Complete the browser-based MCP console so MCP tools, resources, and prompts can be inspected without relying only on terminal commands.",
                "Improve UI polish and navigation for daily portfolio use.",
                "Expand market-data handling, including clearer unsupported-asset messages and provider configuration guidance.",
                "Prepare an IBKR-first synchronization design while keeping manual Fidelity import available.",
                "Add authentication and deployment planning for future internet access.",
                "Develop the submitted article using the prototype as a case study in MCP architecture for financial applications.",
            ],
            styles,
        )
    )

    story.append(pp("Conclusion", styles["h1"]))
    story.append(pp("Trading Journal is a feasible and meaningful project because it combines a practical user-facing application with a current technical research topic. The preliminary prototype already proves the central idea: a portfolio journal can expose useful capabilities through MCP, consume market-data capabilities through another MCP server, and be demonstrated by an external MCP client. The proposed next phase will mature this prototype into a stronger application and a well-supported article that discusses MCP's role, benefits, limitations, and security considerations in the financial domain.", styles["body"]))

    story.append(pp("References", styles["h1"]))
    refs = [
        "[1] Model Context Protocol. (2026). What is the Model Context Protocol (MCP)? https://modelcontextprotocol.io/docs/getting-started/intro",
        "[2] Model Context Protocol. (2025). Specification, 2025-06-18. https://modelcontextprotocol.io/specification/2025-06-18",
        "[3] Soria Parra, D. (2026). The 2026 MCP Roadmap. Model Context Protocol Blog. https://blog.modelcontextprotocol.io/posts/2026-mcp-roadmap/",
        "[4] Eleventh Hour Enthusiast. (2025). Model Context Protocol (MCP): Landscape, Security Threats, and Future Research Directions. https://medium.com/@EleventhHourEnthusiast/model-context-protocol-mcp-landscape-security-threats-and-future-research-directions-488b8d2eade8",
        "[5] Decode the Future. (2026). What Is MCP? Model Context Protocol Explained for 2026. https://decodethefuture.org/en/what-is-mcp-model-context-protocol/",
        "[6] FastMCP Documentation. Getting Started. https://gofastmcp.com/getting-started/welcome",
        "[7] FastAPI Documentation. https://fastapi.tiangolo.com/",
        "[8] SQLAlchemy Documentation. https://docs.sqlalchemy.org/",
        "[9] Alpaca Documentation. Getting Started. https://docs.alpaca.markets/us/docs/getting-started",
        "[10] Alpaca Documentation. About Market Data API. https://docs.alpaca.markets/us/docs/about-market-data-api",
        "[11] Alpaca Market Data Overview. https://alpaca.markets/data",
    ]
    for ref in refs:
        story.append(pp(ref, styles["ref"]))

    def footer(canvas, doc_template):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawCentredString(4.25 * inch, 0.42 * inch, f"Trading Journal MCP Proposal - Page {doc_template.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def render_pdf_for_review() -> None:
    render_dir = OUTPUT_DIR / "pdf_rendered"
    render_dir.mkdir(parents=True, exist_ok=True)
    for old_render in render_dir.glob("page-*.png"):
        old_render.unlink()
    contact_sheet = render_dir / "contact_sheet.png"
    if contact_sheet.exists():
        contact_sheet.unlink()
    font_cache = ROOT / "tmp" / "fontconfig_cache"
    font_cache.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env["XDG_CACHE_HOME"] = str(font_cache)
    env["HOME"] = str(ROOT)
    subprocess.run(
        [
            "/Users/deepaksingla/.cache/codex-runtimes/codex-primary-runtime/dependencies/native/poppler/bin/pdftoppm",
            "-png",
            str(PDF_PATH),
            str(render_dir / "page"),
        ],
        check=True,
        capture_output=True,
        text=True,
        env=env,
    )


def main() -> None:
    diagrams = make_diagrams()
    build_docx(diagrams)
    try:
        render_docx()
    except Exception as exc:
        print(f"DOCX render skipped: {exc}")
    build_pdf(diagrams)
    render_pdf_for_review()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
